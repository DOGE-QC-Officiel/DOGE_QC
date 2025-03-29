import csv
import os
import sys
from docx import Document
from docx.shared import Inches

#################################################################
#                   DATA PREPARATION SECTION
#################################################################
AGF_PROJECTS_FILENAME = 'grants_and_funds.csv'
AGF_projects_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'All Can Funds and Grants', AGF_PROJECTS_FILENAME)

#################################################################
#                DATA ANALYSIS FUNCTIONS SECTION
#################################################################

''' NOTHING '''

#################################################################
#                   UTILITY FUNCTIONS SECTION
#################################################################

def getProjectsAsDictFromCSV(file_path):
    projects = []
    with open(file_path, newline='', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            projects.append(row)

    return projects

def getFundingOrganisationList(project_list):
    organisations = {}
    for project in project_list:
        funding_org = project['owner_org_en']
        organisations[hash(funding_org)] = funding_org
    return organisations;

def getProjectsSince(project_list, year):
    proj_list = []
    for project in project_list:
        proj_year = project['year']
        if int(proj_year) >= int(year):
            proj_list.append(project)
            
    return proj_list

def sumProjectListAgreement(project_list):
    sum_projs = 0
    for project in project_list:
        sum_projs += float(project['agreement_value'])
    return sum_projs

def getLowestValueForColumn(project_list, column_name, type_p):
    lowest_value = None
    for project in project_list:
        try:
            if lowest_value is None or type_p(project[column_name]) < type_p(lowest_value):
                lowest_value = project['year']
        except Exception:
            continue;
    return lowest_value;

def getHighestValueForColumn(project_list, column_name, type_p):
    highest_value = None
    for project in project_list:
        try:
            if highest_value is None or type_p(project[column_name]) > type_p(highest_value):
                highest_value = project[column_name]
        except Exception:
            continue;
    return highest_value;
        

#################################################################
#                  CUSTOM CLASSES FOR DATA HANDLING
#################################################################
''' NOTHING '''

#################################################################
#                  MAIN SCRIPT SECTION
#################################################################
def generateDocxForProjectList(project_list, result_file_path, title):
    document = Document()
    
    # titre du projet
    document.add_heading(title, level=2).bold = True
    document.add_paragraph('')

    for project in project_list:
        # organisme
        proj_p = document.add_paragraph('')
        proj_p.add_run('Proj.#: ').bold = True
        proj_p.add_run(project['NoProj']).add_break()

        proj_p.add_run(' Date: ').bold = True
        proj_p.add_run(project['PeriodeDate']).add_break()
        
        proj_p.add_run(' Nom: ').bold = True
        proj_p.add_run(project['NomProjet']).add_break()

        proj_p.add_run(' Organisme(qc): ').bold = True
        proj_p.add_run(project['NomOrgaQue']).add_break()

        proj_p.add_run(' Pays: ').bold = True
        proj_p.add_run(project['NomPays']).add_break()

        proj_p.add_run(' Montant: ').bold = True
        proj_p.add_run(project['MontantSubvention']).add_break()

        proj_p.add_run(' Statut: ').bold = True
        proj_p.add_run(project['Statut']).add_break()
        
    document.save(result_file_path)

if __name__ == '__main__':
    AGF_data = getProjectsAsDictFromCSV(AGF_projects_file)
    print("Number of Projects :" + str(len(AGF_data)))

    funding_organisations = getFundingOrganisationList(AGF_data)
    print("Number of Funding organisations :" + str(len(funding_organisations)))

    sum_by_org = {}
    
    for org in funding_organisations:
        sum_by_org[funding_organisations[org]] = 0

    for project in AGF_data:
        nom_orga = project['owner_org_en']
        sum_by_org[nom_orga] = sum_by_org[nom_orga] + float(project['agreement_value'])

    as_list = []
    for sum_p in sum_by_org:
        as_list.append({'org':sum_p, 'total':sum_by_org[sum_p]})

    as_list.sort(key=lambda x: x['total'], reverse=True)
    for sum_p in as_list:
        print("Org :" + sum_p['org'] + " | Total:" + str(sum_p['total']))
        

    
