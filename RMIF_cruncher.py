import csv
import os
from docx import Document
from docx.shared import Inches

#################################################################
#                   DATA PREPARATION SECTION
#################################################################
MRIF_PROJECTS_FILENAME = 'mrif_projects.csv'
mrif_projects_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'MRIF', MRIF_PROJECTS_FILENAME)

#################################################################
#                DATA ANALYSIS FUNCTIONS SECTION
#################################################################

''' NOTHING '''

#################################################################
#                   UTILITY FUNCTIONS SECTION
#################################################################

def getMRIFProjectsAsDictFromCSV(file_path):
    mrif_projects = []
    with open(file_path, newline='', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            mrif_projects.append(row)

    return mrif_projects

def getOrganismHashMapFromProjectsDict(project_list):
    org_list = {}
    
    for project in project_list:
        org_list[project['NoOrgaQue']] = {'NomOrgaQue':project['NomOrgaQue'], 'NomOrgaEtran':project['NomOrgaEtran']}

    return org_list
    
def getProjectsSince(project_list, year):
    proj_list = []
    for project in project_list:
        date = project['PeriodeDate']
        start_year = date.split('-')[0]
        if int(start_year) >= int(year):
            proj_list.append(project)
            
    return proj_list
        

#################################################################
#                  CUSTOM CLASSES FOR DATA HANDLING
#################################################################
''' NOTHING '''

#################################################################
#                  MAIN SCRIPT SECTION
#################################################################
def generateDocxForProjectList(project_list, result_file_path, title):
    document = Document()
    
    # titre du projet
    document.add_heading(title, level=2).bold = True
    document.add_paragraph('')

    for project in project_list:
        # organisme
        proj_p = document.add_paragraph('')
        proj_p.add_run('Proj.#: ').bold = True
        proj_p.add_run(project['NoProj']).add_break()

        proj_p.add_run(' Date: ').bold = True
        proj_p.add_run(project['PeriodeDate']).add_break()
        
        proj_p.add_run(' Nom: ').bold = True
        proj_p.add_run(project['NomProjet']).add_break()

        proj_p.add_run(' Organisme(qc): ').bold = True
        proj_p.add_run(project['NomOrgaQue']).add_break()

        proj_p.add_run(' Pays: ').bold = True
        proj_p.add_run(project['NomPays']).add_break()

        proj_p.add_run(' Montant: ').bold = True
        proj_p.add_run(project['MontantSubvention']).add_break()

        proj_p.add_run(' Statut: ').bold = True
        proj_p.add_run(project['Statut']).add_break()
        
    document.save(result_file_path)

if __name__ == '__main__':
    mrif_data = getMRIFProjectsAsDictFromCSV(mrif_projects_file)
    
    '''
    print(len(mrif_data))

    sum_transactions = 0;
    for project in mrif_data:
        sum_transactions += (float(project['Versement1']) + float(project['Versement2']) + float(project['Versement3']) + float(project['Versement4']) + float(project['Versement5']))

    sum_engagements = 0
    for project in mrif_data:
        sum_engagements += float(project['MontantSubvention'])

    all_orgs = getOrganismHashMapFromProjectsDict(mrif_data)
                            
    print("Total international fund sent : " + str(sum_transactions))
    print("Total in engagements : " + str(sum_engagements))
    print("Number of organisations : " + str(len(all_orgs)))

    for org in all_orgs:
        print(all_orgs[org])
    '''
    
    list_since_2018 = getProjectsSince(mrif_data, "2018")
    #generateDocxForProjectList(list_since_2018, os.path.dirname(os.path.abspath(__file__)) + '/mrif_projects_since_2018.docx', 'mrif_projects_since_2018')

    sum_transactions = 0;
    for project in list_since_2018:
        sum_transactions += (float(project['Versement1']) + float(project['Versement2']) + float(project['Versement3']) + float(project['Versement4']) + float(project['Versement5']))

    sum_engagements = 0
    for project in list_since_2018:
        sum_engagements += float(project['MontantSubvention'])
                            
    print("Total international fund sent : " + str(sum_transactions))
    print("Total in engagements : " + str(sum_engagements))

    orgs = getOrganismHashMapFromProjectsDict(mrif_data)

    projs_in_leger = []
    sum_leger = 0;
    keyword = "equitas"
    for project in mrif_data:
        if keyword in project['NomOrgaQue'].lower():
            projs_in_leger.append(project)
            sum_leger += (float(project['Versement1']) + float(project['Versement2']) + float(project['Versement3']) + float(project['Versement4']) + float(project['Versement5']))

    generateDocxForProjectList(projs_in_leger, os.path.dirname(os.path.abspath(__file__)) + '/' + keyword + '.docx', keyword)
    print("Sum leger :" + str(sum_leger))
