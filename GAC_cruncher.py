import os
from xml.etree.ElementTree import ElementTree
import xml.etree.ElementTree as ET
import copy
from docx import Document
from docx.shared import Inches
from pathvalidate import sanitize_filename

#################################################################
#                   DATA PREPARATION SECTION louwy
#################################################################

#represent all active projects in ukraine
#ukraine_tree = ET.parse(os.path.dirname(os.path.abspath(__file__)) + '/data/actif/Ukraine_actives.xml')
#ukraine_projects = ukraine_tree.getroot()

#represent all terminating projects since 01/01/2018
terminating_tree = ET.parse(os.path.dirname(os.path.abspath(__file__)) + '/data/GAC/terminating/all_terminating_from_2018_01_01.xml')
terminating_projects = terminating_tree.getroot()

#represent all active projects ( ukraine included )
active_tree = ET.parse(os.path.dirname(os.path.abspath(__file__)) + '/data/GAC/actif/all_actives.xml')
active_projects = active_tree.getroot()

#represent all closed projects since 01/01/2018
closed_tree = ET.parse(os.path.dirname(os.path.abspath(__file__)) + '/data/GAC/closed/all_closed_since_2018.xml')
closed_projects = closed_tree.getroot()

#################################################################
#                DATA ANALYSIS FUNCTIONS SECTION
#################################################################

#get the country contribution %
def getCountryContributionPercentageNormalized(project, country_contrib_str):

    if country_contrib_str == '':
        return 1.0
    
    countries = project.find('countries')
    contrib_str = ''
    
    for country in countries:
        if country_contrib_str.lower() in country.text.lower():
            contrib_str = country.text
            break

    if contrib_str == '':
        return 1.0 #100%
    else:
        contrib_perc_str = contrib_str.split(' ')
        contrib_perc_str = contrib_perc_str[len(contrib_perc_str)-1]
        contrib_perc_str = contrib_perc_str[:len(contrib_perc_str)-1]
        contrib_perc_str = contrib_perc_str.replace(',', '.')
        return float(contrib_perc_str) / 100.0
        
    

#sums all contribution for a given project list. Will ignore promises and only calculate "sent" transactions
def sumProjectsContributions(projects_list, country_contrib_str):
    proj_sum = 0;
    for project in projects_list.findall('project'):
        contrib_mult = getCountryContributionPercentageNormalized(project, country_contrib_str)
        for transaction in project.find('transactions').findall('transaction'):
            if transaction.get('transactionType') == "Déboursé" or transaction.get('transactionType') == "Disbursement":
                proj_sum += float(transaction.text) * contrib_mult
    return proj_sum

#search in a given project list by keyword. Can include or exluce "0$ money sent" projects
def findProjectsByKeyWord(projects_list, keyword, include_0_projects):
    found_projects = ET.Element("filtered_by_word_"+keyword)

    for project in projects_list.iter('project'):
        to_sum = ET.Element("to_sum")
        to_sum.append(project)
        
        if not include_0_projects and sumProjectsContributions(to_sum, '') == 0:
            continue;

        if keyword == '':
            found_projects.append(project)
            continue
        
        if keyword.lower() in project.find('description').text:
            found_projects.append(project)
            continue;
        
        if keyword.lower() in project.find('title').text:
            found_projects.append(project)
            continue;

    return found_projects

#sort the given project list by highest contribution engagement
def getProjectsSortedByHighestContribution(project_list):
    sorted_projects = ET.Element("sorted_by_highest_contribution")
    for project in sorted(project_list, key=lambda x: float(x.find('maximumContribution').text), reverse=True):
        sorted_projects.append(copy.deepcopy(project))
    return sorted_projects

#sort the given project list by lowest contribution engagement
def getProjectsSortedByLowestContribution(project_list):
    sorted_projects = ET.Element("sorted_by_lowest_contribution")
    for project in sorted(project_list, key=lambda x: float(x.find('maximumContribution').text), reverse=False):
        sorted_projects.append(copy.deepcopy(project))
    return sorted_projects

#return a list of transaction information for given project list
#can ignore or include contribution engagement
#can ignore or include 0$ transactions
def getAllTransactions(project_list, include_promises, include_0):
    transactions = []
    for project in project_list.findall('project'):
        for transaction in project.find('transactions').findall('transaction'):
            if include_promises is True and transaction.get('transactionType') == "Engagement":
                transactions.append({"transaction":transaction, "program_name": project.find('programName').text, "exec_partner": project.find('executingAgencyPartner').text})
            elif transaction.get('transactionType') == "Déboursé":
                if include_0:
                    transactions.append({"transaction":transaction, "program_name": project.find('programName').text, "exec_partner": project.find('executingAgencyPartner').text})
                elif transaction.text is not None and transaction.text != "" and transaction.text != "0":
                    transactions.append({"transaction":transaction, "amount":transaction.text, "program_name": project.find('programName').text, "exec_partner": project.find('executingAgencyPartner').text})
    return transactions

#return the number of total transactions for given project list
#can ignore or include contribution engagement
#can ignore or include 0$ transactions
def nbTransactions(project_list, include_promises, include_0):
    return len(getAllTransactions(project_list, include_promises, include_0))

#################################################################
#                   UTILITY FUNCTIONS SECTION
#################################################################
#returns the xml element as a readable string
def getXMLString(xml_element):
    return ET.tostring(xml_element, encoding='unicode')

def getXMLFromProjectList(project_obj_list, root_name="filtered_projects"):
    root = ET.Element(root_name)
    for proj in project_obj_list:
        root.append(proj.xml_)
    return root

#print some project information 
def outputProjectDescription(project, country_contrib_str, output_func):
    xml_root = ET.Element("to_sum")
    xml_root.append(project)

    output_func("****************************")
    output_func("Titre du programme : " + project.find('title').text)
    output_func("Montant debourser :" + str(sumProjectsContributions(xml_root, country_contrib_str)))
    output_func("Agence partenaire executant :" + project.find('executingAgencyPartner').text)
    output_func()
    output_func(project.find('description').text)
    output_func(getXMLString(project.find('countries')))
    output_func('****************************')

#output all given projects description
def outputAllProjectsDescriptions(project_list, country_contrib_str, output_func):
    for project in project_list.findall('project'):
        outputProjectDescription(project, output_func)
        
#utility function to return a statistic object for given project list
def getProjectsStats(project_list, country_contrib_str):
    stats = {}
    stats['nb_projects'] = str(len(project_list.findall('project')))
    stats['total_contrib'] = str(sumProjectsContributions(project_list, country_contrib_str))
    stats['nb_transactions'] = str(nbTransactions(project_list, True, False))
    stats['nb_transactions_sent'] = str(nbTransactions(project_list, False, False))
    stats['transactions'] = getAllTransactions(project_list, False, False)
    return stats

#output a transaction information
def outputTransaction(transaction, output_func):
    desc_str = "Date:" + transaction['transaction'].get("transactionDate")
    desc_str = desc_str + " | Amount:" + transaction["amount"]
    desc_str = desc_str + " | Program:" + transaction['program_name']
    desc_str = desc_str + " | Executive partner:" + transaction['exec_partner']
    output_func(desc_str)

#output all transactions information
def outputTransactions(transactions, output_func):
    for transaction in transactions:
        outputTransaction(transaction, output_func)

#takes a list of xml root with 'project' elements to merge them into a single root
def mergeXMLProjects(list_of_xml_roots):
    xml_root = ET.Element("merged_projects")
    for xml_projects in list_of_xml_roots:
        for project in xml_projects.findall('project'):
            xml_root.append(project)
    return xml_root

#returns a hash table of the project list and ensures there is no project duplicate
def hashProjectsRemoveDuplicates(xml_project_list):
    project_hash_table = {}
    for project in xml_project_list.findall('project'):
        project_hash_table[hash(getXMLString(project))] = project
    return project_hash_table

def getXMLRootProjectListFromHashTable(project_hash_table):
    xml_root = ET.Element("no_duplicate_projects")
    for hash_val in project_hash_table:
        xml_root.append(project_hash_table[hash_val])
    return xml_root

def removeProjectDuplicates(xml_project_list_root):
    return getXMLRootProjectListFromHashTable(hashProjectsRemoveDuplicates(xml_project_list_root))

def findProjectsByCountries(project_list, country_str):
    xml_root = ET.Element("found_by_country")
    for project in project_list:
        proj_countries = project.find('countries')
        for proj_country in proj_countries:
            if country_str.lower() in proj_country.text.lower():
                xml_root.append(project)

    return xml_root

def ensureDirectoryExists(directory_name):
    os.makedirs(directory_name, exist_ok=True)

def writeStatsToFile(stats, dirname, filename, keyword):
    directory_name = os.path.join(os.path.dirname(os.path.abspath(__file__)),"Search results", dirname)
    output_filename = os.path.join(directory_name, filename)
    ensureDirectoryExists(directory_name)
    
    with open(output_filename, 'w') as f:
        f.write("### Search results for '_" + keyword + "_'<br />\n  ")
        f.write("__Number of projects__ : " + stats['nb_projects'] + "<br />\n")
        f.write("__Total aid sent in $__ :" + str(stats['total_contrib'])  + "<br />\n" )
        f.write("__Number of transactions ( *engagements included* )__ :" + str(stats['nb_transactions'])  + "<br />\n" )
        f.write("__Number of transactions ( *engagements excluded* )__ :" + str(stats['nb_transactions_sent'])  + "<br />\n")
        f.close()
        
    print("Stats written to file : " + output_filename)

def writeXMLProjectListToFile(project_list, dirname, filename):
    directory_name = os.path.join(os.path.dirname(os.path.abspath(__file__)),"Search results", dirname)
    output_filename = os.path.join(directory_name, filename)
    ensureDirectoryExists(directory_name)
    
    with open(output_filename, 'w') as f:
        f.write(getXMLString(project_list))
        
    print("Projects written to file : " + output_filename)
    
def generateDocxForProject(GAC_proj, result_file_path):
    document = Document()
    
    # titre du projet
    document.add_heading(GAC_proj.title_, level=2).bold = True
    document.add_paragraph('')
    
    # organisme
    organism_p = document.add_paragraph('')
    organism_p.add_run('Organisme : ').bold = True
    organism_p.add_run('Affaires Mondiales Canada')
    #p.add_run('italic.').italic = True
    
    # numero de projet
    num_proj_p = document.add_paragraph('')
    num_proj_p.add_run('Numero de projet : ').bold = True
    num_proj_p.add_run(GAC_proj.project_number_)
    
    # Lieu
    location_p = document.add_paragraph('')
    location_p.add_run('Lieu : ').bold = True
    location_p.add_run(", ".join(GAC_proj.regions_))
    
    # agence executive
    exec_partner_p = document.add_paragraph('')
    exec_partner_p.add_run('Agence executive partenaire : ').bold = True
    exec_partner_p.add_run(GAC_proj.exec_agency_partner_)
    
    # type de financement
    fund_type_p = document.add_paragraph('')
    fund_type_p.add_run('Type de financement : ').bold = True
    fund_type_p.add_run(GAC_proj.fund_type_)
    
    # dates
    date_p = document.add_paragraph('')
    date_p.add_run('Dates : ').bold = True
    date_p.add_run(GAC_proj.dates_['start'] + ' au ' + GAC_proj.dates_['end'])
    
    # engagement initial
    engagement_p = document.add_paragraph('')
    engagement_p.add_run('Engagement : ').bold = True
    engagement_p.add_run(GAC_proj.max_contrib_)
    
    # Somme total envoyer
    total_sent_p = document.add_paragraph('')
    total_sent_p.add_run('Total envoye en $ : ').bold = True
    total_sent_p.add_run(str(GAC_proj.getTransactionsTotal()))
    
    # description
    description_p = document.add_paragraph('')
    description_p.add_run('Description : ').bold = True
    description_p.add_run(GAC_proj.description_)
    
    
    document.add_heading('Transactions', level=2)
    
    for transaction in GAC_proj.transactions_:
        # transaction
        description_p = document.add_paragraph('')
        description_p.add_run('Date : ').bold = True
        description_p.add_run(transaction['date'])
        description_p.add_run('Type : ').bold = True
        description_p.add_run(transaction['type'])
        description_p.add_run(' Montant : ').bold = True
        description_p.add_run(transaction['amount'])
   
    document.add_page_break()
    document.save(result_file_path + "/" + sanitize_filename(GAC_proj.title_ + ".docx"))




#################################################################
#                  CUSTOM CLASSES FOR DATA HANDLING
#################################################################

# take a xml tag and return it's inner HTML as text
def getValueFromXML(xml, tag_str):
    return xml.find(tag_str).text

# take a xml and return a child list from a parent and a child tag
# ex : <xml> <parent_tag> <ct1></ct1> <ct2></ct2> </parent_tag> </xml>
def getListFromXML(xml, parent_tag, child_tag):
    return xml.find(parent_tag).findall(child_tag)

# get a xml country and return and object containing the information
def getSplittedCountryValuesFromXML(country_xml):
    splitted_str = country_xml.text.split(' ')
    contrib_perc_str = splitted_str
    contrib_perc_str = splitted_str[len(splitted_str)-1]
    contrib_perc_str = contrib_perc_str[:len(contrib_perc_str)-1]
    contrib_perc_str = contrib_perc_str.replace(',', '.')
    
    return {
        'country' : country_xml.text[:len(country_xml)- len(contrib_perc_str)-2],
        'contribution' : float(contrib_perc_str) / 100.0
    }


# get a xml sector and return an object containing the information
def getSplittedSectorValuesFromXML(sector_xml):
    splitted_str = sector_xml.text.split(' ')
    contrib_perc_str = splitted_str
    contrib_perc_str = splitted_str[len(splitted_str)-1]
    contrib_perc_str = contrib_perc_str[:len(contrib_perc_str)-1]
    contrib_perc_str = contrib_perc_str.replace(',', '.')
    
    return {
        'sector' : sector_xml.text[:len(sector_xml)- len(contrib_perc_str)-2],
        'contribution' : float(contrib_perc_str) / 100.0
    }

# take a xml policy marker and return an object containing the information
def getPolicyMarkerFromXML(policy_marker_xml):
    return {
        'policy' : policy_marker_xml.text,
        'level' : policy_marker_xml.get('level')
    }

# take a xml location and return an object containing the information
def getLocationFromXML(location_xml):
    srs_name = location_xml.get('SRSName')
    srs_splitted = srs_name.split('/')
    epsg = srs_splitted[len(srs_splitted)-3]
    code = srs_splitted[len(srs_splitted)-1]
    lat = location_xml.text.split(' ')[0] if type(location_xml.text) is not type(None) else "0"
    long = location_xml.text.split(' ')[1] if type(location_xml.text) is not type(None) and len(location_xml.text.split(' ')) > 1 else "0"
    return {
        'srs' : epsg+':'+code,
        'lat' : lat,
        'long' : long,
    }

# take a xml transaction and return an object containing the information
def getTransactionFromXML(transaction_xml):
    return {
        'type' : transaction_xml.get('transactionType'),
        'tied_status' : transaction_xml.get('tiedStatus'),
        'amount' : transaction_xml.text,
        'date' : transaction_xml.get('transactionDate')
    }
    
'''
Represent a Global Affairs Canada project
'''
class GAC_Project:
    def __init__(self, project_xml):
        # hold the originL XML for this project. Should not be altered.
        self.xml_ = project_xml
        
        self.project_number_ = getValueFromXML(self.xml_, 'projectNumber')
        self.title_ = getValueFromXML(self.xml_, 'title')
        self.description_ = getValueFromXML(self.xml_, 'description')

        # can be active, terminating or closed
        self.status_ = getValueFromXML(self.xml_, 'status')

        # start and end date of the engagement. A project should be closed after the end date.
        self.dates_ = {
            'start' : getValueFromXML(self.xml_, 'start'),
            'end' : getValueFromXML(self.xml_, 'end')
        }
        
        self.fund_type_ = getValueFromXML(self.xml_, 'financeType')

        # list of countries receiving the money. May be empty
        self.countries_ = [
            getSplittedCountryValuesFromXML(country_xml) for country_xml in getListFromXML(self.xml_, 'countries', 'country')
        ]

        # project executive agency partner. Basically the organisation we give the money to
        self.exec_agency_partner_ = getValueFromXML(self.xml_, 'executingAgencyPartner')

        # type of aid
        self.sectors_ = [
            getSplittedSectorValuesFromXML(sector_xml) for sector_xml in getListFromXML(self.xml_, 'DACSectors', 'DACSectors')
        ]

        # max contribution for the project. Transaction should not exceed this amount ..
        self.max_contrib_ = getValueFromXML(self.xml_, 'maximumContribution')

        # some information about the results of the project
        self.expected_results_ = getValueFromXML(self.xml_, 'expectedResults')
        self.achieved_results_ = getValueFromXML(self.xml_, 'resultsAchieved')

        self.program_name_ = getValueFromXML(self.xml_, 'programName')

        # policy markers seems to define organisation policy compliance in order to receive the fund
        # things like sex equality, biodiversity, climat change etc ...
        # might not be this, but it seems like it
        self.policy_markers_ = [
            getPolicyMarkerFromXML(policy_marker_xml) for policy_marker_xml in getListFromXML(self.xml_, 'policyMarkers', 'policyMarker')
        ]

        # define the region the aid is destined to. Not a particular country. Could be like 'subsahara region' , 'North america' etc..
        self.regions_ = [
            region_xml.text for region_xml in getListFromXML(self.xml_, 'regions', 'region')
        ]

        # some gps information about the project. Number of locations seems to match the number of region...
        self.locations_ = [
            getLocationFromXML(location_xml) for location_xml in getListFromXML(self.xml_, 'Locations', 'location')
        ]

        # all the transaction for this project
        # transaction can either be an engagement or a payment
        self.transactions_ = [
            getTransactionFromXML(transaction_xml) for transaction_xml in getListFromXML(self.xml_, 'transactions', 'transaction')
        ]
      
    def getTransactionsTotal(self):
        sum_transactions = 0.0

        for transaction in self.transactions_:
            if transaction['type'] == "Déboursé" or transaction['type'] == "Disbursement":
                sum_transactions += float(transaction['amount'])
        
        return sum_transactions

'''
    Hold all the GAC_Project in a list and stat them
'''
class GAC_ProjectList:
    def __init__(self, project_list_xml):
        self.xml_ = project_list_xml
        self.projects_ = [GAC_Project(project) for project in project_list_xml.findall("project")]
        self.stats_ = getProjectsStats(project_list_xml, '')

    def findByKeyword(self, keyword, include_0_projects=False):
        found_projects = []

        for project in self.projects_:
            if not include_0_projects and project.getTransactionsTotal() == 0:
                continue
            if keyword == '':
                found_projects.append(project)
                continue
            if keyword.lower() in project.description_.lower():
                found_projects.append(project)
                continue
            if keyword.lower() in project.title_.lower():
                found_projects.append(project)
                continue
        
        return found_projects



#################################################################
#                  MAIN SCRIPT SECTION
#################################################################

def generateDocxForFolder(folder_name):
    with open(os.path.dirname(os.path.abspath(__file__)) + '/' + folder_name + '/projects.xml', 'r') as xml_file:
        project_list_et = ET.parse(xml_file)
        project_list = GAC_ProjectList(project_list_et.getroot())

        for project in project_list.projects_:
            result_file_path = os.path.dirname(os.path.abspath(__file__)) + '/' + folder_name
            generateDocxForProject(project, result_file_path)
    

if __name__ == '__main__':
    '''
    all_projects = mergeXMLProjects([active_projects, closed_projects, terminating_projects])
    all_projects_no_duplicate = removeProjectDuplicates(all_projects)

    gac_proj_list = GAC_ProjectList(all_projects_no_duplicate)
    print(gac_proj_list.stats_['total_contrib'])
    
    COUNTRY = ''
    KEYWORD = 'aga khan'

    all_found_projects = gac_proj_list.findByKeyword(KEYWORD, include_0_projects=False)
    all_found_projects_xml = getXMLFromProjectList(all_found_projects)
    stats = getProjectsStats(all_found_projects, COUNTRY)

    directory_name = KEYWORD
    writeStatsToFile(stats, directory_name, "README.md", KEYWORD)
    writeXMLProjectListToFile(all_found_projects, directory_name, "projects.xml")
    
    generateDocxForFolder("Search results/aga Khan")
    generateDocxForFolder("Search results/bangsamoro")
    generateDocxForFolder("Search results/banque asiatique")
    generateDocxForFolder("Search results/blackrock")    generateDocxForFolder("Search results/covid")
    generateDocxForFolder("Search results/gavi")
    generateDocxForFolder("Search results/LGBT")
    generateDocxForFolder("Search results/parlementaires")
    generateDocxForFolder("Search results/ukraine")
    generateDocxForFolder("Search results/usaid")
    generateDocxForFolder("Search results/vaccin")
    generateDocxForFolder("Search results/irak")
    generateDocxForFolder("Search results/Ukraine")
    '''

    all_projects = mergeXMLProjects([active_projects, closed_projects, terminating_projects])
    all_projects_no_duplicate = removeProjectDuplicates(all_projects)

    gac_proj_list = GAC_ProjectList(all_projects_no_duplicate)
    
    COUNTRY = 'Ukraine'
    KEYWORD = 'Ukraine'

    all_found_projects = gac_proj_list.findByKeyword(KEYWORD, include_0_projects=False)
    stats = getProjectsStats(all_found_projects, COUNTRY)

    directory_name = KEYWORD
    writeStatsToFile(stats, directory_name, "README.md", KEYWORD)
    writeXMLProjectListToFile(all_found_projects, directory_name, "projects.xml")
    generateDocxForFolder("Search results/" + KEYWORD)
    
    
